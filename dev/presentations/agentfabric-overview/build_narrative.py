#!/usr/bin/env python3
"""Build the AgentFabric overview narrative member with python-docx.

Structure authority: narrative-specification.md. Factual authority: source-notes.md.
The narrative and the deck must not diverge on a shared claim.
"""

from __future__ import annotations

import json
import os
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

HERE = Path(__file__).resolve().parent
SOURCE = Path(os.environ.get("AGENTFABRIC_DECK_SOURCE", str(HERE)))
REPO = Path(os.environ.get("AGENTFABRIC_REPO", str(SOURCE.parents[2])))


def _obj_dir() -> Path:
    return Path(os.environ.get("OBJ_DIR") or (REPO / "_build"))


OUT = Path(os.environ.get("AGENTFABRIC_NARRATIVE_OUTPUT") or (SOURCE / "agentfabric-overview.docx"))
PPTX = Path(os.environ.get("AGENTFABRIC_DECK_OUTPUT") or (SOURCE / "agentfabric-overview.pptx"))
INK = RGBColor(0x10, 0x13, 0x17)
STEEL = RGBColor(0x65, 0x70, 0x7C)


def _set_run_font(run, *, name: str, size: Pt, color=INK, bold=False, italic=False):
    run.font.name = name
    run.font.size = size
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:cs"), name)


def heading(doc: Document, level: int, text: str) -> None:
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        _set_run_font(run, name="Calibri", size=Pt(22 - (level * 2)), color=INK, bold=True)


def body(doc: Document, text: str, *, italic: bool = False) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = paragraph.add_run(text)
    _set_run_font(run, name="Calibri", size=Pt(12), color=INK, italic=italic)


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(item)
        _set_run_font(run, name="Calibri", size=Pt(11), color=INK)


def code(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.left_indent = Inches(0.25)
    run = paragraph.add_run(text)
    _set_run_font(run, name="Courier New", size=Pt(9), color=STEEL)


def write_manifest(pptx_path: Path, docx_path: Path) -> Path:
    manifest = {
        "schema": "agentfabric/document-pair-manifest@1",
        "members": {
            "presentation": {
                "local_artifact": str(pptx_path),
                "published_location": None,
                "publication_authorized": False,
            },
            "narrative": {
                "local_artifact": str(docx_path),
                "published_location": None,
                "publication_authorized": False,
            },
        },
        "authoring_package": {
            "root": str(SOURCE),
            "elements": {
                "deck_specification": "deck-specification.md",
                "narrative_specification": "narrative-specification.md",
                "factual_ledger": "source-notes.md",
                "generation_prompts": "prompts",
                "deck_build_source": "build_deck.py",
                "narrative_build_source": "build_narrative.py",
                "regeneration_entry_point": "regenerate.sh",
                "deliverable_links": "current-deliverables.md",
                "qa_record": "qa-ledger.md",
            },
        },
    }
    path = _obj_dir() / "agentfabric-overview" / "document-pair-manifest.json"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(manifest, indent=2) + "\n", encoding="utf-8")
    return path


def build() -> Path:  # noqa: PLR0915 - one linear document, written in order
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    core = doc.core_properties
    core.title = "AgentFabric: one control plane for a fleet of AI agents"
    core.author = "AgentFabric maintainers"
    core.subject = "Narrative member of the AgentFabric overview document pair"

    # ---------------------------------------------------------------- part 1
    heading(doc, 1, "AgentFabric: one control plane for a fleet of AI agents")
    body(
        doc,
        "This is the narrative member of the AgentFabric overview. It accompanies the "
        "20-slide presentation built from deck-specification.md, carries the same claims "
        "in prose at greater depth, and shares one factual ledger with it: "
        "source-notes.md. Where the two members disagree on a shared claim, one of them "
        "is wrong and must be corrected rather than reconciled by the reader.",
    )
    body(
        doc,
        "The reader this is written for is highly technical and completely unfamiliar with "
        "both AgentFabric and the requirements of running many agents and many models at "
        "once. Nothing here assumes prior product knowledge. AgentFabric is the "
        "presentation name for the control plane whose repository, CLI, and modules are "
        "named mac; they are the same system.",
    )

    heading(doc, 2, "What problem this solves")
    body(
        doc,
        "One engineer supervising one coding agent works. The human reads the diff, "
        "notices the wrong turn, and re-steers. Supervision is the human, and it is "
        "genuinely sufficient at that scale. The interesting failure appears when the same "
        "pattern is multiplied: several dozen agents, several models, several machine "
        "classes, and work that outlives any single conversation.",
    )
    body(
        doc,
        "At that point the missing piece is not a better prompt or a smarter agent. It is "
        "that a conversation is ephemeral state. A transcript has no owner, no ordering, "
        "no dependency edges, no lease, and no receipt. Two agents can start the same "
        "work; an agent can stop mid-flight and take the only record of what it was doing "
        "with it; the money spent on model calls cannot be attributed to the work it paid "
        "for; and nobody can answer, afterwards, who approved a change. Those are not "
        "tooling gaps, they are consequences of using a chat log as a system of record.",
    )
    body(
        doc,
        "AgentFabric's position is that a fleet of agents needs the same thing any other "
        "distributed system needs: durable state with ownership, ordering, and evidence. "
        "The unit of truth is a task in a ledger, not a message in a transcript. Agents "
        "become participants in a control plane rather than the control plane itself.",
    )

    heading(doc, 2, "Who this document is for, and what it does not claim")
    body(
        doc,
        "This narrative is for an engineer evaluating adoption, an architect reviewing the "
        "control model, or a manager who wants the deck's claims traced to their source. "
        "It states no return-on-investment figure, no productivity multiple, and no "
        "throughput projection, because none of those are measured. Counted figures "
        "describe surface area and carry the command and the date that produced them. "
        "Where a capability is decided but not yet running, or proposed but not yet "
        "decided, it is labelled as such in the same sentence that describes it.",
    )

    # ---------------------------------------------------------------- part 2
    heading(doc, 1, "What the system does")
    heading(doc, 2, "Ask once, and the fabric carries it to production")
    body(
        doc,
        "The externally visible behaviour is a single pipeline: a request becomes a "
        "durable task; the task is leased to an agent that can actually do it; the agent "
        "executes inside a sandbox; a different agent reviews the result; and publication "
        "leaves provenance behind. Ask, durable task, lease and dispatch, execute, "
        "independent review, publish. Every stage is a boundary owned by the control "
        "plane, and nothing advances because a participant asserts it should.",
    )
    heading(doc, 2, "The three consequences worth stating plainly")
    body(
        doc,
        "First, nothing is lost. A task outlives the chat that produced it, the agent that "
        "claimed it, and the machine that ran it, because the task is a row with a state, "
        "an owner, a lease, dependencies, and history rather than a paragraph in a "
        "transcript.",
    )
    body(
        doc,
        "Second, nothing is blind. Timing, model usage, gate decisions, and provenance are "
        'recorded per step, which is what makes the question "where did the day go, and '
        'what did it cost?" answerable at all.',
    )
    body(
        doc,
        "Third, nothing self-approves. The agent that produced the work does not hold the "
        "authority to accept it. Review is a separate invocation under control-plane-owned "
        "gates. This is an authority boundary rather than a sandbox control, and it is the "
        "single property that makes autonomous agent output safe to merge.",
    )
    heading(doc, 2, "One request, end to end")
    body(
        doc,
        "Read the loop as six hand-offs, each of which writes something down. A human "
        "states intent in a chat channel. The gateway agent owns that conversation, its "
        "personality, and its memory, and it records identity - deliberately nothing else, "
        "because it holds no operational authority. The hub converts intent into a ledger "
        "row with a state and an owner, and issues a time-bounded lease. A worker executes "
        "the leased task inside a sandbox, emitting normalized action events and model "
        "route events as it goes. A reviewer, which is a different invocation, returns a "
        "named decision. Publication records provenance and an evidence pointer.",
    )
    body(
        doc,
        "The separation between the second and third hand-off is the core architectural "
        "decision of the system: conversation, personality, memory, and channels belong to "
        "the human-facing runtime, while tasks, leases, routing, reviews, evidence, secret "
        "handles, runtime manifests, rollout state, and audit trails belong to the control "
        "plane. Authority: README.md and docs/authority-boundary.md.",
    )

    # ---------------------------------------------------------------- part 3
    heading(doc, 1, "Built on the stack, not instead of it")
    heading(doc, 2, "Why re-use is the design position")
    body(
        doc,
        "Isolation, observability, inference capacity, container runtimes, scheduling, and "
        "coding agents all already exist, and each is maintained by people for whom it is "
        "a full-time job. Re-implementing any of them would produce a worse version of "
        "something already available, and would spend the project's budget on the parts of "
        "the problem that are already solved. For an engineering reader the interesting "
        "question about AgentFabric is therefore not what was written; it is what was "
        "deliberately not written.",
    )
    body(
        doc,
        "The rule this package follows when describing re-use is that lineage and "
        "integration are stated separately. A project AgentFabric depends on is named as a "
        "dependency; a project AgentFabric learned from or stays compatible with is named "
        "as a reference. Blurring the two would misrepresent both.",
    )

    heading(doc, 2, "Where AgentFabric leverages NVIDIA technology")
    heading(doc, 3, "NVIDIA OpenShell - the execution security boundary")
    body(
        doc,
        "OpenShell is the confinement authority. Filesystem policy is enforced with "
        "Landlock, syscalls are filtered with seccomp, and network egress goes through an "
        "L7 proxy that is deny-by-default; sandbox lifecycle and the collection of "
        "normalized action events are OpenShell's too. AgentFabric authors the policy and "
        "integrates with the sandbox rather than implementing isolation, and the design "
        "goal is exactly one guardrail authority - the policy file - instead of two "
        "competing ones. Authority: src/mac/openshell_service.py, "
        "src/mac/executor_sandbox.py, src/mac/sandbox_egress.py, "
        "src/mac/openshell_collector.py, docs/openshell-sandbox.md, and ADR 0008, which "
        "makes Docker Engine/Moby OpenShell's only container runtime so that there is a "
        "single enforcement path to reason about.",
    )
    heading(doc, 3, "NVIDIA NeMo Relay - optional observability")
    body(
        doc,
        "Request, task, tool, and model activity is mapped into NeMo Relay scopes when the "
        "relay packaging extra is enabled, which pins nemo-relay==0.3.0. Observability is "
        "mapped into an existing system rather than invented, and it is optional: the "
        "control plane's own event stream remains the durable record. Authority: "
        "src/mac/relay_observability.py, pyproject.toml, "
        "docs/openshell-nemo-relay-integration.md.",
    )
    heading(doc, 3, "NVIDIA HGX - bounded elastic capacity")
    body(
        doc,
        "HGX provides provider-session capacity beyond the statically owned fleet. Two "
        "properties matter more than the capacity itself: it is bounded, and onboarding a "
        "session is an explicit operator action that leaves a durable receipt, which keeps "
        "the machine-onboarding trust boundary intact. Read-only status and planning are "
        "separate from an explicit execute step, and provider work never runs on a "
        "dispatcher or HTTP thread. An automatically scaling elastic executor tier beyond "
        "this operator-driven path is a proposal, not a shipped capability. Authority: "
        "docs/hgx-elastic-capacity.md and ADR 0005 (Proposed).",
    )
    heading(doc, 3, "NVIDIA NemoClaw - a reference, not a deployment")
    body(
        doc,
        "NemoClaw is a compatibility and design reference for the conversational agent "
        "boundary. It is named because the boundary AgentFabric draws around a "
        "human-facing runtime was designed to remain compatible with it, and it is "
        "explicitly not presented as the deployed gateway implementation. Authority: "
        "README.md and docs/hermes-boundary.md.",
    )

    heading(doc, 2, "Where AgentFabric leverages open source")
    body(
        doc,
        "This is the one place a list is the honest form, because the content is an "
        "inventory of load-bearing dependencies grouped by the problem each one already "
        "solved.",
    )
    heading(doc, 3, "State")
    bullets(
        doc,
        [
            "PostgreSQL - the authority for fleet deployments; the test suite runs against "
            "Postgres because that is what the fleet runs.",
            "SQLite - the local development and test path only.",
            "Versioned schema migrations - schema changes are migrations with fail-closed "
            "authority, not an append-only helper list (src/mac/schema_migrations.py).",
        ],
    )
    heading(doc, 3, "Service")
    bullets(
        doc,
        [
            "FastAPI and Uvicorn - the HTTP control surface.",
            "Pydantic - request, response, and configuration models.",
            "httpx - outbound HTTP, including agent-to-hub traffic.",
        ],
    )
    heading(doc, 3, "Execution")
    bullets(
        doc,
        [
            "Docker Engine / Moby - the single container runtime under OpenShell (ADR 0008).",
            "Kubernetes - one of two dispatch targets; a single orchestrator folds claim, "
            "launch, and stuck-Job reconciliation (src/mac/k8s/runner.py).",
            "OpenClaw - the conversational runtime under a control-plane-authored sandbox policy.",
            "OpenAI Codex CLI and OpenCode - coding executors on the route ladder, "
            "selected per task rather than fixed.",
        ],
    )
    heading(doc, 3, "Protocol")
    bullets(
        doc,
        [
            "ACP - agent client protocol endpoints served by the hub (src/mac/acp/).",
            "A2A agent cards - published capability descriptions (src/mac/a2a/card.py).",
            "MCP - the MCP server is a client of the same HTTP surface, not a second "
            "implementation (src/mac/mcp_server.py).",
            "OCSF - the event vocabulary for normalized action events produced under "
            "sandboxed execution (src/mac/openshell_collector.py).",
        ],
    )
    body(
        doc,
        "Every entry above is a dependency, not a fork. The control plane owns none of "
        "those problems, and each protocol is an implemented specification rather than an "
        "aspiration - which is what allows a foreign agent implementation to participate "
        "without a bespoke adapter.",
    )

    heading(doc, 2, "What AgentFabric adds anyway")
    body(
        doc,
        "Re-use is the default, which obliges the project to be equally explicit about the "
        "parts that had to be built. These five mechanisms are what make agent output "
        "trustworthy, and none of the projects above provides them.",
    )
    heading(doc, 3, "A durable task ledger")
    body(
        doc,
        "A task is a state machine with an owner, a lease, dependencies, and history, "
        "persisted in Postgres. Twelve states are defined - open, waiting, blocked, "
        "claimed, running, needs_review, needs_input, stopped, reviewing, completed, "
        "failed, cancelled - of which completed, failed, and cancelled are terminal. "
        "Authority: TaskState and TERMINAL_TASK_STATES in src/mac/models.py, ADR 0004.",
    )
    heading(doc, 3, "Named gate decisions")
    body(
        doc,
        "A gate returns a named decision rather than a boolean, so a refusal is "
        "explainable after the fact instead of appearing as a failed run. The review gate "
        "works this way today (ADR 0011, Accepted); generalizing the contract to every "
        "gate in the system is an accepted-in-principle proposal (ADR 0022, Proposed) "
        "rather than a finished refactor.",
    )
    heading(doc, 3, "An ordered coding-route ladder")
    body(
        doc,
        "Coding executors are plural and pluggable, and selection is an ordered, "
        "capability-filtered ladder rather than a per-worker environment preference. The "
        "current order is opencode, pi, claude, codex, cursor. The order is a fleet "
        "contract, and the reason opencode is first is credential durability rather than "
        "model quality - it keeps a long-lived credential file that can be provisioned to "
        "a worker, where an expiring OAuth session cannot. Do not read the ladder as a "
        "quality ranking. Authority: AGENT_PRIORITY in src/mac/coding_agent.py and "
        "docs/coding-route-ladder.md.",
    )
    heading(doc, 3, "Evidence closure")
    body(
        doc,
        "Build, test, review, and publication artifacts are kept as pointers with the task, "
        "so a completed task carries the numbered evidence that justified it: lint and "
        "format output, test output, the push, and the merge request. A code executor "
        "cannot push or open a merge request without passing that gate, and a task whose "
        "tests fail is routed to review with its full output rather than silently retried. "
        "Authority: AGENTS.md (the mac.worker_evidence.v1 manifest) and "
        "src/mac/observability_service.py.",
    )
    heading(doc, 3, "Break-glass recovery")
    body(
        doc,
        "Recovery from a wedged host is a granted, listable, revocable authorization that "
        "carries a reason, exposed as mac task break-glass, break-glass-list, and "
        "break-glass-revoke. This mechanism exists because an undocumented recovery path "
        "becomes undocumented root access: making the grant first-class is what keeps it "
        "auditable. Authority: docs/break-glass-host-recovery.md.",
    )

    heading(doc, 2, "The trust boundary")
    heading(doc, 3, "What runs inside the sandbox")
    body(
        doc,
        "An agent runs with wide latitude inside a narrow box: its process tree is owned "
        "and reaped, read-only and read-write paths are allow-listed through Landlock, "
        "syscalls are filtered, it never runs as root, egress hosts are declared per "
        "project and per task, and secrets are handles resolved at use. The box - not the "
        "agent's good judgement - is what makes autonomous execution acceptable.",
    )
    heading(doc, 3, "What never crosses it")
    body(
        doc,
        "Four things never cross the boundary: an undeclared network destination, a raw "
        "credential value, another project's repository, and an agent's approval of its own "
        "work. The first three are enforced by the sandbox and the secret store - work "
        "references a secret by name and the value is resolved at use, so a transcript, a "
        "log, or an evidence bundle never carries it. The fourth is an authority boundary "
        "the control plane enforces itself. Authority: docs/openshell-sandbox.md, "
        "docs/secrets-management-guide.md, docs/authority-boundary.md.",
    )

    # ---------------------------------------------------------------- part 4
    heading(doc, 1, "Actors, roles, and who is allowed to decide")
    heading(doc, 2, "The cast, and the single-authority rule")
    body(
        doc,
        "Roles are separated so that the ability to do work, the ability to accept work, "
        "and the ability to change the rules never live in the same place. Every actor "
        "holds exactly one kind of authority.",
    )
    heading(doc, 3, "Requester")
    body(doc, "States intent, in natural language, in a channel. Holds no authority beyond that.")
    heading(doc, 3, "Gateway agent")
    body(
        doc,
        "A stock conversational runtime - OpenClaw in the deployed path - running under a "
        "control-plane-authored sandbox policy. Owns personality, memory, and channels; "
        "owns no operational truth.",
    )
    heading(doc, 3, "Hub")
    body(
        doc,
        "The control plane itself: the ledger, leases, routing, reviews, evidence, secret "
        "handles, and audit trails. It is the only durable authority, and it is the only "
        "actor whose state survives everything else.",
    )
    heading(doc, 3, "Dispatcher")
    body(
        doc,
        "Matches ready work to a node that can actually run it, then leases it. "
        "Dispatchability is a function of dependencies, worker capability, lease state, and "
        "project pause - not a single field.",
    )
    heading(doc, 3, "Worker")
    body(
        doc,
        "Runs exactly one leased task as a supervised, sandboxed process tree, and reports "
        "evidence. A worker holds a lease, never an approval.",
    )
    heading(doc, 3, "Coding executor")
    body(
        doc,
        "Whichever coding agent the route ladder selected for this task. It edits the "
        "repository inside the worker's sandbox. Worker and executor are separate actors "
        "on purpose: the supervision is ours, the coding agent is replaceable.",
    )
    heading(doc, 3, "Reviewer agent")
    body(
        doc,
        "A different invocation from the one that produced the work, returning a named "
        'decision. This is the mechanism behind "nothing self-approves".',
    )
    heading(doc, 3, "Operator")
    body(
        doc,
        "The human with rule-changing authority: grants and revokes recovery "
        "authorizations, pauses projects, onboards nodes and provider sessions. Every grant "
        "is recorded with its reason.",
    )
    body(
        doc,
        "Deliberately absent from this list is any actor that can both perform work and accept it.",
    )

    heading(doc, 2, "The life of one task")
    heading(doc, 3, "A state machine, not a status string")
    body(
        doc,
        "The happy path runs open, waiting, claimed, running, needs_review, reviewing, "
        "completed. Blocked, needs_input, and stopped are held states; failed and cancelled "
        "are terminal. A failed task keeps its evidence: it is a record, not a deletion, "
        "which is what allows a later reviewer to distinguish a flaky environment from a "
        "wrong approach.",
    )
    heading(doc, 3, "The four gates")
    body(
        doc,
        "Gate one is dependency: an unfinished dependency holds the task. Gate two is the "
        "lease: one owner, time-bounded, recoverable. Gate three is evidence: build and "
        "test output exist before review is possible. Gate four is acceptance: an "
        "independent, named decision.",
    )
    heading(doc, 3, "Holds are flags, not states")
    body(
        doc,
        "A staged task carries metadata.no_dispatch=true and is released by removing that "
        "key, rather than by moving through a state. Keeping holds out of the state machine "
        'is why "dispatchable" can remain a computed predicate instead of a field somebody '
        "has to remember to update. Authority: AGENTS.md and src/mac/task_lifecycle.py.",
    )
    heading(doc, 3, "Recovery of stranded and stalled work")
    body(
        doc,
        "Leases expire while a real process may or may not still be alive, so recovery is "
        "explicit rather than implied: there are verbs for stranded and stalled work, and "
        "cancelling a task revokes the lease and aborts the running executor rather than "
        "marking a row and hoping.",
    )

    heading(doc, 2, "Coordination without a switchboard")
    heading(doc, 3, "Broadcast for intent, ledger for consequence")
    body(
        doc,
        "Agents coordinate on a shared broadcast bus - the AgentBus - and act on what they "
        "hear. The hub does not sit in the middle of every conversation; it sits underneath "
        "the durable consequences. That is what makes an interruption such as stand down, "
        "abort, pause, resume, or status cheap and immediate across the whole fleet. The "
        "bus is not the system of record, and must never be described as one: anything "
        "heard on the bus that has a durable consequence lands in the ledger. Authority: "
        "src/mac/agentbus_service.py, src/mac/agentbus_control.py.",
    )
    heading(doc, 3, "Operational learning as secret-free memory")
    body(
        doc,
        "Repository-access outcomes are recorded as secret-free memories, and reviewer "
        "routing prefers recent success while temporarily avoiding a recent authentication "
        "failure. Only the credential source name, redacted host, operation, outcome, "
        "classified failure, and remediation are stored - never a credential value or "
        "authenticated URL. The effect is that a failing credential pattern changes future "
        "routing instead of being retried blindly. Authority: "
        "docs/fleet-operational-learning.md.",
    )

    heading(doc, 2, "The fleet, modelled honestly")
    body(
        doc,
        "Node classes differ because the work differs, and the dispatcher matches a task to "
        "a node that can actually run it rather than pretending the fleet is uniform. "
        "Capability, egress, and secrets are declared per node and per task.",
    )
    heading(doc, 3, "macOS hosts")
    body(
        doc,
        "Host installs managed by launchd, by decision rather than omission: Apple "
        "toolchain work stays native because the container story on macOS is not honest "
        "(ADR 0015, Accepted).",
    )
    heading(doc, 3, "Linux nodes")
    body(
        doc,
        "Containerized execution under Docker Engine/Moby. The pairing of a native node "
        "steward with containerized task execution is accepted with implementation deferred "
        "pending fleet measurement (ADR 0012), and the deck says so rather than claiming it "
        "as running.",
    )
    heading(doc, 3, "Kubernetes")
    body(
        doc,
        "One of two dispatch targets. A single orchestrator folds claim, launch, and "
        "stuck-Job reconciliation, so a stuck Job is reconciled by the same component that "
        "created it (src/mac/k8s/runner.py).",
    )
    heading(doc, 3, "HGX provider sessions")
    body(
        doc,
        "Bounded elastic capacity, onboarded by explicit operator receipt, resolved by "
        "immutable session ID rather than display name. Fleet targets always resolve from "
        "~/.mac/fleets.yaml, which is the source of truth after host swaps.",
    )

    heading(doc, 2, "Models and money")
    heading(doc, 3, "The ordered ladder decides who does the work")
    body(
        doc,
        "Multi-model operation is a routing problem with a cost consequence. The ladder is "
        "ordered and filtered by capability and availability, and it is a fleet contract "
        "rather than a per-worker preference - which is the point of the proposal to make "
        "the route search path itself contractual (ADR 0029, Proposed).",
    )
    heading(doc, 3, "Recording happens where routing happens")
    body(
        doc,
        "Every model call produces one route event carrying input and output token counts "
        "and streaming state, attributed to a task, project, or agent. The router is the "
        "only place that sees every call regardless of which agent made it, which is why "
        "that is where recording belongs.",
    )
    body(
        doc,
        "Being precise about the current state matters here, because this is the claim a "
        "deck most easily overstates. Today the router captures the usage the caller "
        "reported; making the router itself the meter, so that a client which does not ask "
        "for usage cannot produce an unmetered call, is a proposal (ADR 0017). That "
        "proposal quantifies the gap rather than hiding it: over the seven days to "
        "2026-08-19, on 28,352 route events, 8,352 routes (29.5 per cent) recorded a null "
        "input-token count, 5,948 were flagged as streamed without usage, and 2,474 were "
        "attributed to nothing.",
    )
    heading(doc, 3, "Pricing at read time")
    body(
        doc,
        "Cost is not stored. Route events are priced at read time against a models catalog, "
        "so a price-table change re-values history instead of invalidating it, and cost per "
        "task, project, and outcome is a query rather than a re-run. Coverage is itself "
        "measured, so a gap is visible rather than averaged away. Authority: "
        "estimate_route_cost() in src/mac/scientific_optimizer.py and "
        "src/mac/models_catalog.py.",
    )

    heading(doc, 2, "Evidence")
    heading(doc, 3, "What is recorded, and at what granularity")
    body(
        doc,
        "Per task: state history, owner, lease, and dependencies. Per step: timing, token "
        "usage, and the exit decision. Per release: a provenance pin and evidence closure. "
        'Action events arrive as a normalized stream from sandboxed execution, so "what did '
        'the agent actually do" is read rather than reconstructed from a transcript.',
    )
    heading(doc, 3, "An index is not an authorization")
    body(
        doc,
        "This distinction is easy to lose in a slide and is worth stating carefully: the "
        "ledger and its pointers form an index of evidence. Acceptance and publication "
        "remain explicit decisions by an authorized actor. Evidence explains what happened; "
        "it does not, by itself, authorize a release.",
    )

    # ---------------------------------------------------------------- part 5
    heading(doc, 1, "Scope, honestly stated")
    heading(doc, 2, "Measured surface area")
    body(
        doc,
        "The following counts describe the audited tree - jordanhubbard/mac at 2976182, "
        "measured 2026-08-30. They are surface area, not maturity, and they must be "
        "re-measured rather than carried forward when this package is regenerated. Each "
        "figure's command is recorded in source-notes.md.",
    )
    code(
        doc,
        "221  control-plane modules       ls src/mac/*.py | wc -l\n"
        "435  HTTP route declarations     grep route decorators under src/mac\n"
        "458  CLI leaf commands           walk mac.cli.build_parser()\n"
        "  4  first-class CLI objects     project, task, agent, admin\n"
        " 12  task states                 TaskState in src/mac/models.py\n"
        "  5  coding routes               AGENT_PRIORITY in src/mac/coding_agent.py\n"
        "  2  dispatch targets            fleet nodes, Kubernetes",
    )
    body(
        doc,
        "The CLI is the object model, and the HTTP API, the Python client, and the MCP "
        "server are clients of that same surface rather than parallel implementations. That "
        "is why the object count, not the command count, is the meaningful figure.",
    )
    heading(doc, 2, "Implemented")
    body(
        doc,
        "The durable ledger with leases and recovery, sandboxed execution with "
        "deny-by-default egress policy, independent review gates, and per-route "
        "observability events priced at read time are implemented and in use.",
    )
    heading(doc, 2, "Decided, not yet runtime")
    body(
        doc,
        "Agent-initiated review scope is accepted (ADR 0016). The pairing of a native node "
        "steward with containerized execution on Linux is accepted with implementation "
        "deferred pending fleet measurement (ADR 0012). This is the category most decks "
        "omit: an accepted decision whose implementation is deferred is neither a shipped "
        "capability nor a proposal, and collapsing it into either one is what makes an "
        "architecture document untrustworthy.",
    )
    heading(doc, 2, "Proposed")
    body(
        doc,
        "Metering enforced at the router (ADR 0017), the route-search-path contract "
        "(ADR 0029), the task view as a graph under progressive disclosure (ADR 0018), and "
        "the retrieval and extraction pipeline (ADR 0030) are proposals. Each ADR's own "
        "status line is the authority for its placement here.",
    )

    # ---------------------------------------------------------------- part 6
    heading(doc, 1, "How to adopt it")
    heading(doc, 2, "Start with one project and one real workload")
    body(
        doc,
        "Register one repository as a project, create real tasks in the ledger, and let the "
        "fleet claim them. One real workload exercises the parts that matter - leases, "
        "gates, evidence, and review - in a way that a demonstration task does not, because "
        "the fabric's value appears exactly when a second agent has to trust the first "
        "one's output.",
    )
    heading(doc, 2, "Grow node classes and coding routes as work demands")
    body(
        doc,
        "Add a node class when the work needs it - a macOS host for Apple toolchain work, a "
        "Linux node or Kubernetes for containerized execution, an HGX session for bounded "
        "extra capacity - and add a coding route when a task class needs a different "
        "executor. Both are declarations, not forks.",
    )
    heading(doc, 2, "Measure before adding capacity")
    body(
        doc,
        "Read cost and throughput from the ledger before buying more of anything. The "
        "reason to put a control plane underneath agents at all is that it can answer that "
        "question; a fleet that cannot be audited is a fleet that cannot be trusted with a "
        "repository, and it also cannot be sized.",
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    return OUT


def main() -> None:
    out = build()
    document = Document(str(out))
    headings = [p for p in document.paragraphs if p.style.name.startswith("Heading")]
    print(f"built narrative -> {out} ({len(headings)} headings)")
    if not PPTX.is_file():
        raise SystemExit(f"presentation artifact missing: {PPTX}")
    manifest = write_manifest(PPTX, out)
    print(f"document-pair manifest -> {manifest}")


if __name__ == "__main__":
    main()
